import docx
newfile=docx.Document("E:/python/1.docx")
newfile.add_paragraph("我是谁")
newfile.add_heading("帅气的我",level=1)
newfile.save("E:/python/1.docx")